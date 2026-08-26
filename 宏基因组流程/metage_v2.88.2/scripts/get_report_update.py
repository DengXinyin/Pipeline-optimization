#!/usr/bin/env python
# -*- coding: utf-8 -*-
# Update version of get_report.py

import os
import sys
import time
import json
import glob
import locale
import argparse
import logging
import subprocess

import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] %(levelname)s: %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
log = logging.getLogger(__name__)


def table_replace(docx, value, value_re):
    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if value in run.text:
                            run.text = run.text.replace(value, value_re)


def replace_in_runs(paragraph, old, new):
    """在段落的 runs 中替换文本（尽量保持原有样式）。"""
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)


def apply_display_name_map(docx, display_name_map):
    """把 docx 中所有 internal_id 替换为 display_name。"""
    if not display_name_map:
        return
    for old, new in display_name_map.items():
        # 替换表格
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_runs(paragraph, old, new)
        # 替换普通段落
        for paragraph in docx.paragraphs:
            replace_in_runs(paragraph, old, new)


def RP_img(img_ls, tag, paragraphs):
    try:
        imgs = glob.glob(img_ls)
        try:
            locale.setlocale(locale.LC_ALL, '')
            imgs = sorted(imgs, key=locale.strxfrm)
        except Exception:
            imgs = sorted(imgs)
        if not imgs:
            log.warning('[%s] 未找到图片: %s', tag, img_ls)
            return
        pic = imgs[0]
        for i in range(len(paragraphs)):
            if tag in paragraphs[i].text:
                if i + 1 >= len(paragraphs) or not paragraphs[i + 1].runs:
                    log.warning('[%s] 模板段落结构不符，无法插入图片', tag)
                    return
                run = paragraphs[i + 1].runs[0]
                run.clear()
                run.add_picture(pic, width=Cm(10))
                log.info('[%s] 插入图片: %s', tag, pic)
                break
    except Exception as e:
        log.error('[%s] 插入图片失败: %s', tag, e)


def get_businfo(datadir):
    sample = pd.read_csv(os.path.join(datadir, 'sample-metadata.tsv'), sep='\t', skiprows=[1])
    project_info_path = os.path.join(datadir, 'project_info.json')
    if not os.path.exists(project_info_path):
        raise FileNotFoundError(f'project_info.json 不存在: {project_info_path}')
    with open(project_info_path, encoding='utf-8') as f:
        bus_info = json.load(f)
    name = bus_info['客户名称']
    partment = bus_info['客户单位']
    NO = bus_info['项目编号']
    TI = str(time.strftime('%Y-%m-%d', time.localtime()))
    return sample, name, partment, NO, TI


def _is_key_figure(img_ls, tag):
    """判断是否为关键图（只插入overview类图，不插每个feature的统计检验图）。"""
    low = img_ls.lower()
    # 排除各类 per-feature 差异分析图
    if 'statistical_test_analysis' in low:
        return False
    if 'taxstatistical_analysis' in low:
        return False
    if 'functionstatistical_analysis' in low:
        return False
    if '/6.' in low and ('anosim' in low or 'adonis' in low or 'mrpp' in low):
        return False
    # 排除 lefse
    if '/9.lefse' in low or 'lefse' in low:
        return False
    return True


def Micro_RP(datadir, resdir, docx_path, analyse, binning, image_mode='full', display_name_map=None):
    if analyse == 'yes' and binning == 'yes':
        template = os.path.join(docx_path, 'metage_v2.88.2_bins.docx')
    elif analyse == 'yes' and binning == 'no':
        template = os.path.join(docx_path, 'metage_v2.88.2.docx')
    elif analyse == 'no':
        template = os.path.join(docx_path, 'metage_v2.88.2_Noana.docx')
    else:
        raise ValueError(f'无效参数 analyse={analyse}, binning={binning}')

    if not os.path.exists(template):
        raise FileNotFoundError(f'报告模板不存在: {template}')

    micro_docx = Document(template)
    paragraphs = micro_docx.paragraphs
    sample, name, partment, NO, TI = get_businfo(datadir)

    table_replace(micro_docx, '姓名', name)
    table_replace(micro_docx, '检测部', partment)
    table_replace(micro_docx, '目号', NO)
    table_replace(micro_docx, '时间', TI)

    image_tasks = [
        (os.path.join(resdir, 'Result/group*/1-data_quality/*/error_rate.png'), '原始数据碱基质量值分布图'),
        (os.path.join(resdir, 'Result/group*/1-data_quality/*/ATGC_content.png'), '碱基含量分布图'),
        (os.path.join(resdir, 'Result/group*/1-data_quality/*/reads_quality_summary.png'), '原始数据组成图'),
        (os.path.join(resdir, 'Result/group*/2-Assembly/contig_length.png'), 'Contigs长度分布'),
        (os.path.join(resdir, 'Result/group*/3-GenePredict/gene_length.png'), 'gene catalogue长度分布图'),
        (os.path.join(resdir, 'Result/group*/4-GeneAbundance/Sample_correlation/sample.corr_heatmap.png'), '样品间相关系数热图'),
        (os.path.join(resdir, 'Result/group*/4-GeneAbundance/Venn/upset.png'), '韦恩图(Venn Graph)或者upset图'),
        (os.path.join(resdir, 'Result/group*/5-TaxAnnotation/3.Barplot/Samples/All/*.png'), '多样本群落结构柱状图示例图'),
        (os.path.join(resdir, 'Result/group*/5-TaxAnnotation/4.Bar_tree/All/*.png'), '聚类树柱状图组合示例'),
        (os.path.join(resdir, 'Result/group*/5-TaxAnnotation/5.Heatmap/Samples/All/*.png'), '物种组成聚类热图如下'),
        (os.path.join(resdir, 'Result/group*/5-TaxAnnotation/6.Beta_diversity_analysis/*/*/1.PCA/*.png'), 'PCA图如下所示'),
        (os.path.join(resdir, 'Result/group*/5-TaxAnnotation/6.Beta_diversity_analysis/*/*/2.PCoA/*.png'), 'PCoA图如下所示'),
        (os.path.join(resdir, 'Result/group*/5-TaxAnnotation/6.Beta_diversity_analysis/*/*/3.NMDS/*.png'), 'NMDS分析结果图如下'),
        (os.path.join(resdir, 'Result/group*/5-TaxAnnotation/7.alpha_diversity_analysis/*/*.png'), '样本组间alpha多样性指数箱体图如下'),
        (os.path.join(resdir, 'Result/group*/6-TaxStatistical_analysis/*/*/1.ANOVA/*.png'), '组间Anova方差分析柱状图'),
        (os.path.join(resdir, 'Result/group*/6-TaxStatistical_analysis/*/*/2.wilcoxon/*.png'), '组间秩和检验差异物种柱状图'),
        (os.path.join(resdir, 'Result/group*/6-TaxStatistical_analysis/*/*/3.Stamp/*.png'), '组间Stamp分析物种柱状图'),
        (os.path.join(resdir, 'Result/group*/6-TaxStatistical_analysis/*/*/4.Random_Forest/*.png'), '随机森林分析图如下所示'),
        (os.path.join(resdir, 'Result/group*/6-TaxStatistical_analysis/*/*/5.metagenomeSeq/*.png'), 'MetagenomeSeq差异热图如下'),
        (os.path.join(resdir, 'Result/group*/6-TaxStatistical_analysis/*/*/6.Anosim/*.png'), 'Anosim结果如图'),
        (os.path.join(resdir, 'Result/group*/6-TaxStatistical_analysis/*/*/9.Lefse/*.png'), 'LDA值分布柱状图'),
        (os.path.join(resdir, 'Result/group*/7-FunctionAnnotation/*/1.Barplot/*.png'), '功能丰度柱形图如下'),
        (os.path.join(resdir, 'Result/group*/7-FunctionAnnotation/*/2.Heatmap/*.png'), '功能丰度热图如下'),
        (os.path.join(resdir, 'Result/group*/7-FunctionAnnotation/*/3.PCA/*.png'), '基于功能丰度的PCA分析'),
        (os.path.join(resdir, 'Result/group*/7-FunctionAnnotation/*/4.PCoA/*.png'), '基于功能丰度的PCoA分析'),
        (os.path.join(resdir, 'Result/group*/7-FunctionAnnotation/*/5.NMDS/*.png'), '基于功能丰度的NMDS分析'),
        (os.path.join(resdir, 'Result/group*/8-FunctionStatistical_analysis/*/1.ANOVA/*.png'), '基于功能丰度的ANOVA分析差异柱状图'),
        (os.path.join(resdir, 'Result/group*/8-FunctionStatistical_analysis/*/2.wilcoxon/*.png'), '基于功能丰度的秩和检验差异柱状图'),
        (os.path.join(resdir, 'Result/group*/8-FunctionStatistical_analysis/*/3.Stamp/*.png'), '基于功能丰度的stamp差异图'),
        (os.path.join(resdir, 'Result/group*/8-FunctionStatistical_analysis/*/4.Random_Forest/*.png'), '功能通路的随机森林分析图'),
        (os.path.join(resdir, 'Result/group*/8-FunctionStatistical_analysis/*/5.metagenomeSeq/*.png'), '显著差异功能热图'),
        (os.path.join(resdir, 'Result/group*/8-FunctionStatistical_analysis/*/9.Lefse/*.png'), '基于功能丰度的LDA柱状图'),
        (os.path.join(resdir, 'Result/group*/9-METABOLIC/*/1.Barplot/*.png'), '不同循环功能丰度柱形图'),
        (os.path.join(resdir, 'Result/group*/9-METABOLIC/*/2.Heatmap/*.png'), '不同循环功能丰度热图'),
        (os.path.join(resdir, 'Result/group*/9-METABOLIC/*/3.PCA/*.png'), '不同循环的PCA分析如下图'),
        (os.path.join(resdir, 'Result/group*/9-METABOLIC/*/4.PCoA/*.png'), '不同循环的PCoA分析如下'),
        (os.path.join(resdir, 'Result/group*/9-METABOLIC/*/5.NMDS/*.png'), '不同循环的NMDS分析如下'),
        (os.path.join(resdir, 'Result/group*/9-METABOLIC/*/6.Statistical_test_analysis/1.ANOVA/*.png'), '不同循环的ANOVA分析差异柱状图'),
        (os.path.join(resdir, 'Result/group*/9-METABOLIC/*/6.Statistical_test_analysis/2.wilcoxon/*.png'), '不同循环的秩和检验差异柱状图'),
        (os.path.join(resdir, 'Result/group*/9-METABOLIC/*/6.Statistical_test_analysis/3.Stamp/*.png'), '不同循环的stamp差异图'),
        (os.path.join(resdir, 'Result/group*/9-METABOLIC/*/6.Statistical_test_analysis/4.Random_Forest/*.png'), '不同循环的随机森林分析图'),
        (os.path.join(resdir, 'Result/group*/9-METABOLIC/*/6.Statistical_test_analysis/5.metagenomeSeq/*.png'), '不同循环的差异功能热图'),
        (os.path.join(resdir, 'Result/group*/9-METABOLIC/*/6.Statistical_test_analysis/9.Lefse/*.png'), '不同循环的LDA柱状图'),
        (os.path.join(resdir, 'Result/group*/10-ARG/1.Barplot/*.png'), 'ARG丰度柱形图如下'),
        (os.path.join(resdir, 'Result/group*/10-ARG/2.Heatmap/*.png'), 'ARG功能丰度聚类热图'),
        (os.path.join(resdir, 'Result/group*/10-ARG/3.PCA/*.png'), 'ARG功能丰度PCA分析图'),
        (os.path.join(resdir, 'Result/group*/10-ARG/4.PCoA/*.png'), 'ARG功能丰度PCoA分析图如下'),
        (os.path.join(resdir, 'Result/group*/10-ARG/5.NMDS/*.png'), 'ARG功能丰度NMDS分析图'),
        (os.path.join(resdir, 'Result/group*/10-ARG/6.Statistical_test_analysis/1.ANOVA/*.png'), 'ARG的ANOVA分析差异功能柱状图'),
        (os.path.join(resdir, 'Result/group*/10-ARG/6.Statistical_test_analysis/2.wilcoxon/*.png'), 'ARG的秩和检验差异功能柱状图'),
        (os.path.join(resdir, 'Result/group*/10-ARG/6.Statistical_test_analysis/3.Stamp/*.png'), 'ARG差异检验柱状图如下'),
        (os.path.join(resdir, 'Result/group*/10-ARG/6.Statistical_test_analysis/4.Random_Forest/*.png'), 'ARG随机森林分析'),
        (os.path.join(resdir, 'Result/group*/10-ARG/6.Statistical_test_analysis/5.metagenomeSeq/*.png'), 'ARG差异功能热图'),
        (os.path.join(resdir, 'Result/group*/10-ARG/6.Statistical_test_analysis/9.Lefse/*.png'), 'ARG的LDA值柱状图'),
        (os.path.join(resdir, 'Result/group*/11-VFDB/1.Barplot/*.png'), '毒力因子基因丰度柱形图'),
        (os.path.join(resdir, 'Result/group*/11-VFDB/2.Heatmap/*.png'), '毒力因子基因聚类热图'),
        (os.path.join(resdir, 'Result/group*/11-VFDB/3.PCA/*.png'), '毒力因子基因PCA分析图'),
        (os.path.join(resdir, 'Result/group*/11-VFDB/4.PCoA/*.png'), '毒力因子基因PCoA分析'),
        (os.path.join(resdir, 'Result/group*/11-VFDB/5.NMDS/*.png'), '毒力因子基因NMDS分析'),
        (os.path.join(resdir, 'Result/group*/11-VFDB/6.Statistical_test_analysis/1.ANOVA/*.png'), '毒力因子基因ANOVA分析'),
        (os.path.join(resdir, 'Result/group*/11-VFDB/6.Statistical_test_analysis/2.wilcoxon/*.png'), '毒力因子基因秩和检验'),
        (os.path.join(resdir, 'Result/group*/11-VFDB/6.Statistical_test_analysis/3.Stamp/*.png'), '毒力因子基因差异检验柱状图'),
        (os.path.join(resdir, 'Result/group*/11-VFDB/6.Statistical_test_analysis/4.Random_Forest/*.png'), '毒力因子基因随机森林分析'),
        (os.path.join(resdir, 'Result/group*/11-VFDB/6.Statistical_test_analysis/5.metagenomeSeq/*.png'), '毒力因子基因差异功能热图'),
        (os.path.join(resdir, 'Result/group*/11-VFDB/6.Statistical_test_analysis/9.Lefse/*.png'), '毒力因子基因LDA值柱状图'),
        (os.path.join(resdir, 'Result/group*/12-mobileOG/1.Barplot/*.png'), '可移动基因元件丰度柱形图如下：'),
        (os.path.join(resdir, 'Result/group*/12-mobileOG/2.Heatmap/*.png'), '可移动基因元件聚类热图：'),
        (os.path.join(resdir, 'Result/group*/12-mobileOG/3.PCA/*.png'), '可移动基因元件PCA分析图：'),
        (os.path.join(resdir, 'Result/group*/12-mobileOG/4.PCoA/*.png'), '可移动基因元件PCoA分析：'),
        (os.path.join(resdir, 'Result/group*/12-mobileOG/5.NMDS/*.png'), '可移动基因元件NMDS分析：'),
        (os.path.join(resdir, 'Result/group*/12-mobileOG/6.Statistical_test_analysis/1.ANOVA/*.png'), '可移动基因元件ANOVA分析：'),
        (os.path.join(resdir, 'Result/group*/12-mobileOG/6.Statistical_test_analysis/2.wilcoxon/*.png'), '可移动基因元件秩和检验：'),
        (os.path.join(resdir, 'Result/group*/12-mobileOG/6.Statistical_test_analysis/3.Stamp/*.png'), '可移动基因元件差异检验柱状图如下所示：'),
        (os.path.join(resdir, 'Result/group*/12-mobileOG/6.Statistical_test_analysis/4.Random_Forest/*.png'), '可移动基因元件随机森林分析（排名前10）：'),
        (os.path.join(resdir, 'Result/group*/12-mobileOG/6.Statistical_test_analysis/5.metagenomeSeq/*.png'), '可移动基因元件差异功能热图如下：'),
        (os.path.join(resdir, 'Result/group*/12-mobileOG/6.Statistical_test_analysis/9.Lefse/*.png'), '可移动基因元件LDA值柱状图如下所示：'),
        (os.path.join(resdir, 'Result/group*/13-BacMet2/1.Barplot/*.png'), '重金属抗性基因丰度柱形图如下：'),
        (os.path.join(resdir, 'Result/group*/13-BacMet2/2.Heatmap/*.png'), '重金属抗性基因聚类热图：'),
        (os.path.join(resdir, 'Result/group*/13-BacMet2/3.PCA/*.png'), '重金属抗性基因PCA分析图：'),
        (os.path.join(resdir, 'Result/group*/13-BacMet2/4.PCoA/*.png'), '重金属抗性基因PCoA分析：'),
        (os.path.join(resdir, 'Result/group*/13-BacMet2/5.NMDS/*.png'), '重金属抗性基因NMDS分析：'),
        (os.path.join(resdir, 'Result/group*/13-BacMet2/6.Statistical_test_analysis/1.ANOVA/*.png'), '重金属抗性基因ANOVA分析：'),
        (os.path.join(resdir, 'Result/group*/13-BacMet2/6.Statistical_test_analysis/2.wilcoxon/*.png'), '重金属抗性基因秩和检验：'),
        (os.path.join(resdir, 'Result/group*/13-BacMet2/6.Statistical_test_analysis/3.Stamp/*.png'), '重金属抗性基因差异检验柱状图如下所示：'),
        (os.path.join(resdir, 'Result/group*/13-BacMet2/6.Statistical_test_analysis/4.Random_Forest/*.png'), '重金属抗性基因随机森林分析（排名前10）：'),
        (os.path.join(resdir, 'Result/group*/13-BacMet2/6.Statistical_test_analysis/5.metagenomeSeq/*.png'), '重金属抗性基因差异功能热图如下：'),
        (os.path.join(resdir, 'Result/group*/13-BacMet2/6.Statistical_test_analysis/9.Lefse/*.png'), '重金属抗性基因LDA值柱状图如下所示：'),
        (os.path.join(resdir, 'Result/group*/14-QS/1.Barplot/*.png'), '群体感应基因丰度柱形图如下：'),
        (os.path.join(resdir, 'Result/group*/14-QS/2.Heatmap/*.png'), '群体感应基因聚类热图：'),
        (os.path.join(resdir, 'Result/group*/14-QS/3.PCA/*.png'), '群体感应基因PCA分析图：'),
        (os.path.join(resdir, 'Result/group*/14-QS/4.PCoA/*.png'), '群体感应基因PCoA分析：'),
        (os.path.join(resdir, 'Result/group*/14-QS/5.NMDS/*.png'), '群体感应基因NMDS分析：'),
        (os.path.join(resdir, 'Result/group*/14-QS/6.Statistical_test_analysis/1.ANOVA/*.png'), '群体感应基因ANOVA分析：'),
        (os.path.join(resdir, 'Result/group*/14-QS/6.Statistical_test_analysis/2.wilcoxon/*.png'), '群体感应基因秩和检验：'),
        (os.path.join(resdir, 'Result/group*/14-QS/6.Statistical_test_analysis/3.Stamp/*.png'), '群体感应基因差异检验柱状图如下所示：'),
        (os.path.join(resdir, 'Result/group*/14-QS/6.Statistical_test_analysis/4.Random_Forest/*.png'), '群体感应基因随机森林分析（排名前10）：'),
        (os.path.join(resdir, 'Result/group*/14-QS/6.Statistical_test_analysis/5.metagenomeSeq/*.png'), '群体感应基因差异功能热图如下：'),
        (os.path.join(resdir, 'Result/group*/14-QS/6.Statistical_test_analysis/9.Lefse/*.png'), '群体感应基因LDA值柱状图如下所示：'),
        # COG 注释
        (os.path.join(resdir, 'Result/group*/15-COG/1.Barplot/*.png'), 'COG功能丰度柱形图如下：'),
        (os.path.join(resdir, 'Result/group*/15-COG/2.Heatmap/*.png'), 'COG功能丰度聚类热图：'),
        (os.path.join(resdir, 'Result/group*/15-COG/3.PCA/*.png'), 'COG功能丰度PCA分析图：'),
        (os.path.join(resdir, 'Result/group*/15-COG/4.PCoA/*.png'), 'COG功能丰度PCoA分析：'),
        (os.path.join(resdir, 'Result/group*/15-COG/5.NMDS/*.png'), 'COG功能丰度NMDS分析：'),
        (os.path.join(resdir, 'Result/group*/15-COG/6.Statistical_test_analysis/1.ANOVA/*.png'), 'COG的ANOVA分析差异功能柱状图：'),
        (os.path.join(resdir, 'Result/group*/15-COG/6.Statistical_test_analysis/2.wilcoxon/*.png'), 'COG的秩和检验差异功能柱状图：'),
        (os.path.join(resdir, 'Result/group*/15-COG/6.Statistical_test_analysis/3.Stamp/*.png'), 'COG差异检验柱状图如下所示：'),
        (os.path.join(resdir, 'Result/group*/15-COG/6.Statistical_test_analysis/4.Random_Forest/*.png'), 'COG随机森林分析（排名前10）：'),
        (os.path.join(resdir, 'Result/group*/15-COG/6.Statistical_test_analysis/9.Lefse/*.png'), 'COG的LDA值柱状图如下所示：'),
        # MetaCyc 代谢通路
        (os.path.join(resdir, 'Result/group*/16-MetaCyc/1.Barplot/*.png'), 'MetaCyc代谢通路丰度柱形图如下：'),
        (os.path.join(resdir, 'Result/group*/16-MetaCyc/2.Heatmap/*.png'), 'MetaCyc代谢通路聚类热图：'),
        (os.path.join(resdir, 'Result/group*/16-MetaCyc/3.PCA/*.png'), 'MetaCyc代谢通路PCA分析图：'),
        (os.path.join(resdir, 'Result/group*/16-MetaCyc/4.PCoA/*.png'), 'MetaCyc代谢通路PCoA分析：'),
        (os.path.join(resdir, 'Result/group*/16-MetaCyc/5.NMDS/*.png'), 'MetaCyc代谢通路NMDS分析：'),
        (os.path.join(resdir, 'Result/group*/16-MetaCyc/6.Statistical_test_analysis/1.ANOVA/*.png'), 'MetaCyc的ANOVA分析差异通路柱状图：'),
        (os.path.join(resdir, 'Result/group*/16-MetaCyc/6.Statistical_test_analysis/2.wilcoxon/*.png'), 'MetaCyc的秩和检验差异通路柱状图：'),
        (os.path.join(resdir, 'Result/group*/16-MetaCyc/6.Statistical_test_analysis/3.Stamp/*.png'), 'MetaCyc差异检验柱状图如下所示：'),
        (os.path.join(resdir, 'Result/group*/16-MetaCyc/6.Statistical_test_analysis/4.Random_Forest/*.png'), 'MetaCyc随机森林分析（排名前10）：'),
        (os.path.join(resdir, 'Result/group*/16-MetaCyc/6.Statistical_test_analysis/9.Lefse/*.png'), 'MetaCyc的LDA值柱状图如下所示：'),
    ]

    if binning == 'yes':
        image_tasks += [
            (os.path.join(resdir, 'Result/binning/2.Bin_Plot/*.png'), '分箱GC_coverage图：'),
            (os.path.join(resdir, 'Result/binning/3.Bin_Abundance/*.png'), '分箱丰度聚类热图：'),
        ]

    if image_mode == 'none':
        log.info('image_mode=none，跳过所有结果图片插入')
        filtered_tasks = []
    elif image_mode == 'key':
        filtered_tasks = [(img_ls, tag) for img_ls, tag in image_tasks if _is_key_figure(img_ls, tag)]
        log.info('image_mode=key，共 %d 个图片任务，保留关键图 %d 个', len(image_tasks), len(filtered_tasks))
    else:
        filtered_tasks = image_tasks

    for img_ls, tag in filtered_tasks:
        RP_img(img_ls, tag, paragraphs)

    # 替换报告中的 internal_id 为 display_name
    if display_name_map:
        log.info('应用 display_name_map 到报告')
        apply_display_name_map(micro_docx, display_name_map)
        append_display_name_map_table(micro_docx, display_name_map)

    report_docx = os.path.join(resdir, 'report.docx')
    micro_docx.save(report_docx)
    log.info('保存报告: %s', report_docx)


def append_display_name_map_table(docx, display_name_map):
    """在报告末尾追加样本名映射表，使报告能直观反映 display_name 变更。"""
    if not display_name_map:
        return

    # 添加标题段落
    title = docx.add_paragraph()
    run = title.add_run('样本名映射表 (Sample Name Mapping)')
    run.bold = True
    run.font.size = Pt(12)

    # 添加说明段落
    desc = docx.add_paragraph(
        '下表列出原始样本名（internal_id）与报告中展示名（display_name）的对应关系。'
    )
    desc.runs[0].font.size = Pt(10)

    # 添加表格（兼容不同环境内置样式）
    table = docx.add_table(rows=1, cols=2)
    try:
        table.style = 'Table Grid'
    except Exception:
        pass
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '原始样本名 (internal_id)'
    hdr_cells[1].text = '展示样本名 (display_name)'

    for iid, display in sorted(display_name_map.items()):
        row_cells = table.add_row().cells
        row_cells[0].text = iid
        row_cells[1].text = display

    log.info('已在报告末尾追加样本名映射表，共 %d 条映射', len(display_name_map))


def load_display_name_map(map_path):
    """读取 display_name_map.tsv，返回 {internal_id: display_name}。"""
    if not map_path or not os.path.exists(map_path):
        return {}
    try:
        df = pd.read_csv(map_path, sep='\t', dtype=str)
        mapping = {}
        for _, r in df.iterrows():
            iid = str(r['internal_id']) if pd.notna(r['internal_id']) else ''
            display = str(r['display_name']) if pd.notna(r['display_name']) else ''
            if iid and display and iid != display:
                mapping[iid] = display
        log.info('加载 display_name_map: %d 条映射', len(mapping))
        return mapping
    except Exception as e:
        log.warning('读取 display_name_map 失败: %s', e)
        return {}


def main():
    parser = argparse.ArgumentParser(description='Generate Result report (update version)')
    parser.add_argument('-I', '--i_datadir', type=str, required=True, help='the dir of sample-metadata.tsv and project_info.json')
    parser.add_argument('--analyse', type=str, choices=['yes', 'no'], required=True, help='Whether to analyze')
    parser.add_argument('--binning', type=str, choices=['yes', 'no'], required=True, help='Whether to binning')
    parser.add_argument('--res_dir', type=str, required=True, help='the dir of Result')
    parser.add_argument('--micro_docx_path', type=str, default='/root/microbiome/microbiome/metage_v2.88.2', help='the dir of micro docx templates')
    parser.add_argument('--image-mode', type=str, choices=['full', 'key', 'none'], default='full',
                        help='图片插入模式：full=全部（默认），key=只插入关键图，none=不插入结果图')
    parser.add_argument('--display-name-map', type=str, default='',
                        help='display_name_map.tsv 路径，用于把报告中的 internal_id 替换为 display_name')
    args = parser.parse_args()

    datadir = os.path.abspath(args.i_datadir)
    res_dir = os.path.abspath(args.res_dir)
    docx_path = os.path.abspath(args.micro_docx_path)

    if not os.path.isdir(datadir):
        log.error('数据目录不存在: %s', datadir)
        sys.exit(1)
    if not os.path.isdir(res_dir):
        log.error('结果目录不存在: %s', res_dir)
        sys.exit(1)

    display_name_map = load_display_name_map(args.display_name_map)

    try:
        Micro_RP(datadir, res_dir, docx_path, args.analyse, args.binning, args.image_mode,
                 display_name_map=display_name_map)
        log.info('开始转换 report.docx -> report.pdf')
        cmd = [
            'libreoffice7.5', '--headless', '--convert-to', 'pdf',
            '--outdir', res_dir, os.path.join(res_dir, 'report.docx')
        ]
        subprocess.run(cmd, check=True)
        log.info('get_report 完成')
    except subprocess.CalledProcessError as e:
        log.error('LibreOffice 转换失败: %s', e)
        sys.exit(1)
    except Exception as e:
        log.error('get_report 运行失败: %s', e)
        sys.exit(1)


if __name__ == '__main__':
    main()
